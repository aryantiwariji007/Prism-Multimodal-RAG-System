# Parse PDF 
from PyPDF2 import PdfReader
from docx import Document
import os
import sys
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# Add the backend directory to the Python path for progress service
backend_dir = Path(__file__).parent.parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

try:
    from app.services.progress_service import progress_service
except ImportError:
    # Fallback if progress service is not available
    progress_service = None

def parse_document(path, file_id=1, progress_file_id=None):
    """
    Parse either PDF or DOCX files based on file extension
    """
    file_extension = os.path.splitext(path)[1].lower()
    chunks = []
    
    if file_extension == '.pdf':
        chunks = parse_pdf(path, file_id, progress_file_id)
    elif file_extension == '.docx':
        chunks = parse_docx(path, file_id, progress_file_id)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Only PDF and DOCX files are supported.")
    
    return chunks

def parse_pdf(path, file_id=1, progress_file_id=None):
    """
    Parse PDF files using pdfplumber with Heuristic Parsing.
    Detects Headers based on font size (Text > 1.1x Median Size).
    """
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Opening PDF (Semantic Mode)...")

    import pdfplumber
    import numpy as np
    from ingestion.chunker import document_chunker

    structured_items = []
    
    try:
        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            
            # 1. Analyze Font Statistics to find "Body Text" size
            all_font_sizes = []
            # Sample first 5 pages for speed
            for p in pdf.pages[:5]:
                for char in p.chars:
                    all_font_sizes.append(char.get("size", 10))
            
            if all_font_sizes:
                median_size = np.median(all_font_sizes)
                header_threshold = median_size * 1.15
            else:
                median_size = 10
                header_threshold = 12
            
            logger.info(f"PDF Parsing: Median Font={median_size}, Header Threshold={header_threshold}")

            for i, page in enumerate(pdf.pages):
                # Extract words/lines with layout info
                # pdfplumber extract_words usually gives x0, top, bottom, etc.
                # We need to reconstruct lines and check their avg font size.
                
                # Extract tables first (to exclude them from text flow if needed, OR treat them as special blocks)
                # For this Agentic Pipeline, let's treat tables as explicit "Table" blocks
                tables = page.find_tables()
                table_bboxes = [t.bbox for t in tables]
                
                # Process Tables
                if tables:
                    for t_idx, table in enumerate(page.extract_tables()):
                        import pandas as pd
                        try:
                            # Filter out None and clean rows
                            cleaned_table = []
                            for row in table:
                                cleaned_table.append([str(cell).strip() if cell else "" for cell in row])
                            
                            df_table = pd.DataFrame(cleaned_table)
                            # Basic Header Detection: promote first row if it has distinct values
                            if len(df_table) > 1:
                                df_table.columns = [str(c) if str(c).strip() else f"Col_{j}" for j, c in enumerate(df_table.iloc[0])]
                                df_table = df_table.drop(df_table.index[0])
                            
                            t_md = df_table.to_markdown(index=False)
                            t_text = f"#### [Page {i+1} | Table {t_idx+1}]\n{t_md}"
                        except Exception as table_err:
                            logger.warning(f"Failed to create MD table: {table_err}")
                            # Fallback to simple join
                            t_text = f"--- Table {t_idx+1} (Page {i+1}) ---\n" + "\n".join([" | ".join([str(c) for c in r if c]) for r in table])
                        
                        if t_text:
                            structured_items.append({
                                "type": "table",
                                "text": t_text,
                                "page": i + 1,
                                "parent_titles": []
                            })

                # Process Text (filtering out table areas to avoid duplication?)
                # Simplification: use extract_text but splitting by newline and checking font size of that line?
                # Hard with 'extract_text' which loses font info. We must iterate layout objects or `.chars`.
                # Better approach for high-precision:
                
                # Get all words with font info
                words = page.extract_words(extra_attrs=["size"])
                
                # Group words into lines based on 'top' coordinate
                lines = {} # top_coord -> list of words
                for w in words:
                    # check if inside a table bbox
                    cx = (w['x0'] + w['x1']) / 2
                    cy = (w['top'] + w['bottom']) / 2
                    in_table = False
                    for bbox in table_bboxes:
                        if bbox[0] < cx < bbox[2] and bbox[1] < cy < bbox[3]:
                            in_table = True
                            break
                    if in_table:
                        continue
                        
                    # Group by rough Y lines (tolerance 3px)
                    found_line = False
                    for y_key in lines.keys():
                        if abs(y_key - w['top']) < 5: # 5px tolerance
                            lines[y_key].append(w)
                            found_line = True
                            break
                    if not found_line:
                        lines[w['top']] = [w]
                
                # Sort lines by Y (top)
                sorted_y = sorted(lines.keys())
                
                for y in sorted_y:
                    line_words = sorted(lines[y], key=lambda x: x['x0'])
                    line_text = " ".join([w['text'] for w in line_words])
                    
                    # Avg font size of line
                    avg_size = sum([w['size'] for w in line_words]) / len(line_words)
                    
                    if avg_size > header_threshold:
                        # It's a header
                        structured_items.append({
                            "type": "heading",
                            "text": line_text,
                            "level": 2 if avg_size < header_threshold * 1.5 else 1,
                            "page": i + 1
                        })
                    else:
                        # Body text
                        structured_items.append({
                            "type": "text",
                            "text": line_text,
                            "page": i + 1
                        })

                # Update progress
                if progress_service and progress_file_id:
                    progress = ((i + 1) / total_pages) * 90
                    progress_service.update_progress(progress_file_id, 1, f"Parsed page {i+1}...", 10 + progress)
    
    except Exception as e:
        logger.error(f"Semantic PDF Parse failed: {e}. Falling back.")
        return _parse_pdf_fallback(path, file_id, progress_file_id)

    # Chunk result
    file_name = os.path.basename(path)
    chunks = document_chunker.chunk_structured_content(structured_items, file_id, file_name=file_name)
    return chunks

def _parse_pdf_fallback(path, file_id, progress_file_id):
    """Original PyPDF2 implementation as fallback"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        chunks = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception as e:
                logger.warning(f"Failed to extract text from page {i} of {path}: {e}")
                text = ""
            chunks.append({"file_id": file_id, "page": i, "text": text})
        return chunks
    except Exception as e:
        logger.error(f"Fallback PDF parsing failed for {path}: {e}")
        return [{"file_id": file_id, "page": 0, "text": ""}]

def parse_docx(path, file_id=1, progress_file_id=None):
    """
    Parse DOCX files using python-docx with Style detection.
    Returns a list of structured elements:
    [{ "type": "heading", "text": "...", "level": 1}, { "type": "text", "text": "..." }]
    """
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Opening DOCX file...")
    
    doc = Document(path)
    structured_items = []
    
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Extracting structured text...", 20)
    
    # Iterate through document elements in order would be ideal, 
    # but python-docx separates tables and paragraphs. 
    # We will approximate by reading paragraphs, and if we hit a table place-holder...
    # Actually, iterating paragraphs is the main flow. Tables are tricky in stream.
    # We will stick to the standard paragraph iteration and handle tables separately or 
    # try to interleave if possible. For now, we process paragraphs then tables is acceptable 
    # BUT we miss the context of WHERE the table is.
    # Better approach: iterating paragraphs and tables in order is hard with python-docx xml.
    # We will stick to: Process Paragraphs (detect Headers) -> Append formatted tables at the end (or as found if we can).
    
    # Actually, let's just do paragraphs respecting styles.
    # Tables often contain data that is referenced by surrounding text.
    # For a v1 "Contextual", parsing paragraphs in order is robust.
    
    iter_elements = doc.iter_inner_content() if hasattr(doc, 'iter_inner_content') else doc.paragraphs # iter_inner_content allows mixed
    
    # If iter_inner_content isn't available (old versions), we might miss table order. 
    # We will use doc.paragraphs for text and append tables at the end as "Appendices" effectively 
    # unless we can map them.
    
    # Check parent titles stack
    parent_titles = []
    
    # Simple formatting: Paragraphs
    total_elements = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        style_name = para.style.name.lower()
        
        # Detect Headings
        if 'heading 1' in style_name:
            structured_items.append({"type": "heading", "text": text, "level": 1, "page": 0, "parent_titles": []})
            parent_titles = [text] # Reset stack
        elif 'heading 2' in style_name:
            structured_items.append({"type": "heading", "text": text, "level": 2, "page": 0, "parent_titles": list(parent_titles)})
            if len(parent_titles) > 0:
                 # sibling or child? simplify: always append if level 2
                 if len(parent_titles) >= 2: parent_titles = parent_titles[:1]
                 parent_titles.append(text)
            else:
                 parent_titles = [text]
        elif 'heading 3' in style_name:
             structured_items.append({"type": "heading", "text": text, "level": 3, "page": 0, "parent_titles": list(parent_titles)})
             # simplistic stack logic
        else:
            # Normal text
            structured_items.append({"type": "text", "text": text, "page": 0, "parent_titles": list(parent_titles)})

    # Tables - Append them effectively as "Table Data"
    for t_idx, table in enumerate(doc.tables):
        import pandas as pd
        try:
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            
            df_table = pd.DataFrame(data)
            if len(df_table) > 1:
                df_table.columns = df_table.iloc[0]
                df_table = df_table.drop(df_table.index[0])
            
            t_md = df_table.to_markdown(index=False)
            t_text = f"#### [DOCX Table {t_idx+1}]\n{t_md}"
        except Exception as e:
            logger.warning(f"DOCX table conversion failed: {e}")
            t_text = f"--- Table {t_idx + 1} ---\n" + "\n".join([" | ".join([c.text for c in r.cells]) for r in table.rows])
        
        if t_text:
            structured_items.append({
                "type": "table", 
                "text": t_text, 
                "page": 0, 
                "parent_titles": ["Table Data"]
            })

    from ingestion.chunker import document_chunker
    # Use the new structured chunker
    file_name = os.path.basename(path)
    chunks = document_chunker.chunk_structured_content(structured_items, file_id, file_name=file_name)

    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "DOCX structure parsing completed", 100)
    
    return chunks
